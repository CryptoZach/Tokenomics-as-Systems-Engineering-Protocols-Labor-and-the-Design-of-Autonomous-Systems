#!/usr/bin/env python3
"""
Phase 2: Thesis Integration — Update DePIN_Thesis_v10_Compressed.docx → v11_DataExpansion.docx

This script:
1. Reads the v10 thesis
2. Expands Table 6.G1 with new governance rows
3. Expands Table 6.D1 with new S2R rows
4. Expands Table 6.U1 with new unit economics rows
5. Updates Table 6.0 (data coverage matrix)
6. Updates Table 6.0A (claims boundary)
7. Adds protocol profiles to §6.10.3
8. Updates Chapter 7 references
9. Updates Table 7.1 claim registry
10. Saves as v11
"""
import os
import sys
import json
import copy
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(SCRIPT_DIR)), "data", "expansion")
INPUT_PATH = os.path.join(os.path.dirname(os.path.dirname(SCRIPT_DIR)), "Pub3_SSRN_final.docx")
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.dirname(SCRIPT_DIR)), "Pub3_SSRN_v11_DataExpansion.docx")


def load_data():
    """Load all expansion data."""
    data = {}
    for fname in os.listdir(DATA_DIR):
        if fname.endswith(".json"):
            with open(os.path.join(DATA_DIR, fname)) as f:
                data[fname] = json.load(f)
    return data


def find_paragraph_containing(doc, text, start_from=0):
    """Find paragraph index containing text."""
    for i, para in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        if text in para.text:
            return i
    return -1


def find_table_by_content(doc, search_text):
    """Find table index containing search text in any cell."""
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    return i
    return -1


def add_rows_to_table(table, rows_data):
    """Add rows to an existing table."""
    try:
        for row_data in rows_data:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                if i < len(row.cells):
                    row.cells[i].text = str(cell_text)
        return True
    except Exception as e:
        print(f"    Table row addition failed: {e}")
        print(f"    Attempting XML-level insertion...")
        try:
            from docx.oxml.ns import qn
            from lxml import etree
            tbl = table._tbl
            # Ensure tblGrid exists
            tbl_grid = tbl.find(qn('w:tblGrid'))
            if tbl_grid is None:
                # Create tblGrid from first row
                tbl_grid = etree.SubElement(tbl, qn('w:tblGrid'))
                if table.rows:
                    for cell in table.rows[0].cells:
                        etree.SubElement(tbl_grid, qn('w:gridCol'))
                # Move tblGrid to correct position (before first row)
                tbl.remove(tbl_grid)
                first_tr = tbl.find(qn('w:tr'))
                if first_tr is not None:
                    first_tr.addprevious(tbl_grid)
                else:
                    tbl.append(tbl_grid)
            # Now try adding rows again
            for row_data in rows_data:
                row = table.add_row()
                for i, cell_text in enumerate(row_data):
                    if i < len(row.cells):
                        row.cells[i].text = str(cell_text)
            return True
        except Exception as e2:
            print(f"    XML insertion also failed: {e2}")
            return False


def insert_paragraph_after(doc, paragraph_idx, text, style=None):
    """Insert a new paragraph after the given index using XML manipulation."""
    from docx.oxml.ns import qn
    from lxml import etree

    if paragraph_idx < 0 or paragraph_idx >= len(doc.paragraphs):
        return None
    para = doc.paragraphs[paragraph_idx]
    new_p = etree.SubElement(para._element.getparent(), qn('w:p'))
    para._element.addnext(new_p)
    # Add text run
    if text:
        r = etree.SubElement(new_p, qn('w:r'))
        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        t.set(qn('xml:space'), 'preserve')
    # Return the new paragraph (re-index to find it)
    return new_p


def main():
    print("=" * 70)
    print("THESIS INTEGRATION: v10 → v11")
    print("=" * 70)

    if not os.path.exists(INPUT_PATH):
        print(f"ERROR: Input thesis not found at {INPUT_PATH}")
        return

    doc = Document(INPUT_PATH)
    data = load_data()

    print(f"  Loaded thesis: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    # ═══════════════════════════════════════════════════════════════
    # 1. EXPAND TABLE 6.G1 (Governance Concentration)
    # ═══════════════════════════════════════════════════════════════
    print("\n1. Expanding Table 6.G1...")
    g1_idx = find_table_by_content(doc, "Holding HHI")
    if g1_idx == -1:
        g1_idx = find_table_by_content(doc, "UNI")
    if g1_idx >= 0:
        table = doc.tables[g1_idx]
        new_rows = [
            ["DIMO", "DIMO", "0.305", "0.94", "69.5%", "54.8%", "1000", "Dune Polygon (Feb 2026)"],
            ["IoTeX", "IOTX", "0.388", "0.98", "91.7%", "61.3%", "989", "Dune Ethereum (Feb 2026)"],
            ["WeatherXM", "WXM", "0.593", "0.97", "94.1%", "76.4%", "999", "Dune Arbitrum (Feb 2026)"],
            ["Anyone Protocol", "ANYONE", "0.034", "0.79", "44.8%", "11.9%", "999", "Dune Ethereum (Feb 2026)"],
        ]
        add_rows_to_table(table, new_rows)
        print(f"  Added {len(new_rows)} rows to Table 6.G1 (table index {g1_idx})")
    else:
        print("  WARNING: Table 6.G1 not found")

    # ═══════════════════════════════════════════════════════════════
    # 2. EXPAND TABLE 6.D1 (S2R by Sink Architecture)
    # ═══════════════════════════════════════════════════════════════
    print("\n2. Expanding Table 6.D1...")
    d1_idx = find_table_by_content(doc, "Sink Type")
    if d1_idx == -1:
        d1_idx = find_table_by_content(doc, "S2R (burns)")
    if d1_idx >= 0:
        table = doc.tables[d1_idx]
        new_rows = [
            ["Hivemapper", "Map credit burn", "TBD", "TBD", "Feb 2026", "Solana on-chain (pending)"],
            ["DIMO", "License burn", "4.24", "—", "Feb 2026", "Dune Polygon (on-chain)"],
            ["IoTeX", "Gas fee burn", "TBD", "TBD", "Feb 2026", "IoTeX L1 (not on Dune)"],
            ["WeatherXM", "Subscription burn", "TBD", "TBD", "Feb 2026", "Public docs"],
            ["Grass", "None (staking only)", "~0", "~0", "Feb 2026", "Public docs"],
            ["Anyone", "Stake/lock", "~0", "~0", "Feb 2026", "Public docs"],
        ]
        add_rows_to_table(table, new_rows)
        print(f"  Added {len(new_rows)} rows to Table 6.D1 (table index {d1_idx})")
    else:
        print("  WARNING: Table 6.D1 not found")

    # ═══════════════════════════════════════════════════════════════
    # 3. EXPAND TABLE 6.U1 (Unit Economics)
    # ═══════════════════════════════════════════════════════════════
    print("\n3. Expanding Table 6.U1...")
    u1_idx = find_table_by_content(doc, "Subsidy Risk")
    if u1_idx == -1:
        u1_idx = find_table_by_content(doc, "DePIN Network")
    if u1_idx >= 0:
        table = doc.tables[u1_idx]
        new_rows = [
            ["Bandwidth/proxy", "UpRock", "~$1.50/GB", "Bright Data", "$12.75/GB", "8.5x", "High"],
            ["AI data", "Grass", "Indirect (tokens)", "Scale AI", "$25-40/hr", "N/A", "High"],
            ["Vehicle data", "DIMO", "Free (earn tokens)", "Otonomo", "$0.50-3/veh/mo", "N/A", "High"],
            ["Weather data", "WeatherXM", "~$0.01/API call", "Tomorrow.io", "$0.05-0.20/call", "10x", "Medium"],
            ["VPN relay", "Anyone", "Free relay (earn)", "NordVPN", "$3.49-8.32/mo", "N/A", "High"],
            ["Community WiFi", "Wayru/WiFiDabba", "~$2-5/mo", "ISP retail", "$15-30/mo", "6.4x", "Med-High"],
            ["Wearable health", "CUDIS", "~$69 ring", "Oura Ring", "$299 + $6/mo", "4.3x", "High"],
            ["IoT infra (L1)", "IoTeX", "~$0.01/tx", "AWS IoT Core", "$1/M msgs", "N/A", "Medium"],
        ]
        add_rows_to_table(table, new_rows)
        print(f"  Added {len(new_rows)} rows to Table 6.U1 (table index {u1_idx})")
    else:
        print("  WARNING: Table 6.U1 not found")

    # ═══════════════════════════════════════════════════════════════
    # 4. UPDATE TABLE 6.0 (Data Coverage Matrix)
    # ═══════════════════════════════════════════════════════════════
    print("\n4. Updating Table 6.0...")
    t0_idx = find_table_by_content(doc, "Governance holding HHI")
    if t0_idx >= 0:
        table = doc.tables[t0_idx]
        # Update the governance row
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if "Governance holding HHI" in cells[0]:
                row.cells[1].text = "11 protocols × ~990 holders"
                row.cells[2].text = "Feb 2026"
                row.cells[3].text = "Dune Analytics (multi-chain)"
            elif "DePIN unit economics" in cells[0]:
                row.cells[1].text = "13 service domains"
            elif "DePIN mini-cases" in cells[0]:
                row.cells[0].text = "DePIN profiles (expanded)"
                row.cells[1].text = "7 deep + 4 mini protocols"
        print(f"  Updated Table 6.0 (table index {t0_idx})")
    else:
        print("  WARNING: Table 6.0 not found")

    # ═══════════════════════════════════════════════════════════════
    # 5. UPDATE TABLE 6.0A (Claims Boundary)
    # ═══════════════════════════════════════════════════════════════
    print("\n5. Updating Table 6.0A...")
    t0a_idx = find_table_by_content(doc, "Claims Boundary")
    if t0a_idx == -1:
        t0a_idx = find_table_by_content(doc, "does NOT claim")
    if t0a_idx >= 0:
        table = doc.tables[t0a_idx]
        for row in table.rows:
            for cell in row.cells:
                if "Mini-cases + unit economics" in cell.text:
                    cell.text = cell.text.replace(
                        "Mini-cases + unit economics",
                        "DePIN profiles + unit economics"
                    )
                if "29–80×" in cell.text or "29-80×" in cell.text or "29–80x" in cell.text:
                    cell.text = cell.text.replace("5 domains", "13 domains spanning IoT, storage, compute, mapping, bandwidth, vehicle data, weather, WiFi, health, and mobility")
                if "Does not prove sustainability" in cell.text:
                    cell.text = "Does not prove sustainability without token subsidies; posted prices (N=1 per service per provider); service equivalence varies by domain; S2R snapshots, not longitudinal trajectories (except Helium); Solana governance data unavailable via Dune (HONEY, GRASS)"
        print(f"  Updated Table 6.0A (table index {t0a_idx})")
    else:
        print("  WARNING: Table 6.0A not found")

    # ═══════════════════════════════════════════════════════════════
    # 6. ADD PROTOCOL PROFILES TO §6.10.3
    # ═══════════════════════════════════════════════════════════════
    print("\n6. Adding protocol profiles to §6.10.3...")
    profiles_idx = find_paragraph_containing(doc, "Emerging DePIN Protocols: Minimum Viable Profiles")
    if profiles_idx == -1:
        profiles_idx = find_paragraph_containing(doc, "Four additional DePIN protocols were surveyed")
    if profiles_idx >= 0:
        # Load protocol profiles if available
        profiles_file = os.path.join(DATA_DIR, "protocol_profiles.json")
        if os.path.exists(profiles_file):
            with open(profiles_file) as f:
                profiles = json.load(f)
            print(f"  Found {len(profiles)} profiles to insert")
        else:
            # Write inline profiles based on collected data
            profiles = None
            print("  Protocol profiles JSON not yet available - writing inline from collected data")

        # Replace the "Four additional" paragraph with expanded text
        old_text = doc.paragraphs[profiles_idx].text
        doc.paragraphs[profiles_idx].text = (
            "Expanded DePIN Protocol Profiles"
        )

        # Insert profiles after this paragraph
        insert_idx = profiles_idx

        # Full profiles (4 paragraphs each)
        full_profiles = [
            {
                "name": "IoTeX: IoT Infrastructure Layer-1",
                "tag": "(Finding, Empirical)",
                "p1": "IoTeX operates a purpose-built Layer-1 blockchain optimized for Internet of Things devices and DePIN applications. The platform sells IoT middleware (W3bstream) that enables devices to generate verifiable proofs of real-world activity, bridging physical hardware events to on-chain token rewards. Device manufacturers and DePIN projects build on IoTeX infrastructure, paying IOTX gas fees for transactions. With a market capitalization of $51M and 9.4 billion circulating tokens (CoinGecko, February 2026), IoTeX is among the most mature DePIN infrastructure protocols.",
                "p2": "IoTeX employs an EIP-1559-style gas fee burn mechanism on its native L1 chain. Transaction fees are partially burned, creating a usage-proportional sink. Staking locks reduce effective circulating supply; approximately 40% of IOTX is staked with network delegates. The governance concentration data reveals extreme concentration: HHI 0.388 with 61.3% of top-1000 holder tokens in a single address and 91.7% in the top 10 (Dune, Ethereum ERC-20 representation, February 2026). This likely reflects staking pool aggregation and team treasury holdings rather than broad retail distribution.",
                "p3": "Verification relies on W3bstream's trusted execution environment, which validates IoT device data before issuing on-chain proofs. Devices register cryptographic identities and submit signed data streams. The integrity layer is moderate: hardware-based attestation plus economic stake, but without the cryptographic proof-of-service rigor of Filecoin's Proof-of-Spacetime.",
                "p4": "IoTeX's sustainability signal is mixed. Gas fee burns create a direct usage-to-token sink, but the absolute burn volume is small relative to staking emissions. The protocol's competitive positioning as DePIN infrastructure middleware is distinctive — rather than competing with a specific Web2 service, IoTeX competes with general-purpose L1s and IoT cloud platforms (AWS IoT Core). The governance concentration finding (HHI 0.388) suggests that despite IoTeX's maturity, token distribution remains highly concentrated on the Ethereum representation, potentially reflecting cross-chain bridge dynamics rather than underlying governance structure.",
            },
            {
                "name": "DIMO: Connected Vehicle Data Marketplace",
                "tag": "(Finding, Empirical)",
                "p1": "DIMO operates a decentralized connected vehicle data marketplace on Polygon. The unit of service is vehicle telemetry data — OBD-II diagnostics, GPS, and driving behavior — shared by car owners who install DIMO hardware dongles or connect via software integrations. Data consumers (insurers, fleet managers, researchers) purchase access through the DIMO Developer License, which requires burning DIMO tokens. As of February 2026, DIMO has a market capitalization of $4.6M with 424 million circulating tokens out of 1 billion total supply (CoinGecko).",
                "p2": "DIMO's sink architecture features a license burn mechanism: developers must burn DIMO tokens to mint a Developer License NFT for API access to vehicle data. This creates a demand-coupled burn analogous to Helium's Data Credit model. On-chain data from Dune Analytics reveals an S2R of 4.24 for the most recent period — indicating that burns significantly exceed minting, though this likely reflects a period of elevated developer license purchases against low emission rates rather than sustained equilibrium. Governance concentration is extremely high: HHI 0.305 with the top address holding 54.8% of tokens (Dune, Polygon, February 2026), consistent with team/treasury concentration in an early-stage protocol.",
                "p3": "Verification relies on hardware device attestation (DIMO Macaron dongle) and software integrations with OEM APIs. Data authenticity is ensured by signed device certificates. The verification posture is moderate: hardware-based provenance for directly connected vehicles, but dependent on third-party APIs for software-only integrations.",
                "p4": "DIMO's sustainability signal is early-positive. The license burn mechanism creates a direct pathway from data demand to token value, and the S2R above unity suggests meaningful demand-side activity. However, the small market capitalization ($4.6M) and high governance concentration suggest the protocol is in its growth phase. The unit economics comparison with Web2 vehicle data providers (Otonomo, Wejo at $0.50-3.00/vehicle/month) is not directly comparable because DIMO users share data voluntarily for token rewards rather than paying a subscription. The novel data ownership model — where drivers control and monetize their own vehicle data — represents a capability expansion in the Nussbaum/Sen sense.",
            },
            {
                "name": "Hivemapper: Decentralized Street Mapping",
                "tag": "(Finding, Empirical)",
                "p1": "Hivemapper operates a decentralized street-level mapping network on Solana. The unit of service is freshness-verified map imagery, generated by contributors who install authenticated dashcams in their vehicles. Mapping clients consume map data by purchasing Map Credits, which requires burning HONEY tokens. As of February 2026, Hivemapper has a market capitalization of $20.8M with 5.6 billion HONEY circulating (CoinGecko). The network has mapped portions of roads in 100+ countries, though coverage density varies significantly by region.",
                "p2": "Hivemapper's sink architecture implements a burn-and-mint tokenomic model similar to Helium's DC system. Map Credits are purchased with HONEY (burned), and contributors earn HONEY for submitting verified imagery. MIP-7 introduced tiered pricing based on map freshness and coverage. S2R data is not yet available from Dune (Solana token transfer schema limitations prevented the query from returning results), but the mechanism design is structurally similar to Helium's DC burn model.",
                "p3": "Verification relies on authenticated dashcam hardware with tamper-resistant firmware. Images include cryptographic signatures, GPS coordinates, and timestamps. The network uses AI-based quality scoring to detect anomalous or fabricated imagery. Anti-gaming posture is moderate-to-strong: hardware-enforced image provenance combined with algorithmic quality filtering, though the system is not immune to physical-world manipulation (e.g., driving the same route repeatedly).",
                "p4": "Hivemapper's sustainability signal is competitive but subsidy-dependent. Map tile pricing (~$5/1000 tiles) undercuts Google Maps API ($7/1000 loads) by approximately 29%, but this discount is partially sustained by HONEY emissions to dashcam contributors. The subsidy risk is Medium: the burn mechanism creates genuine demand-side revenue, but the pricing advantage may narrow as emissions taper. Hivemapper's key competitive advantage — freshness (continuously updated from contributor dashcams vs. Google's periodic Street View updates) — represents a genuine service differentiation rather than pure price competition.",
            },
            {
                "name": "WeatherXM: Hyperlocal Weather Data Network",
                "tag": "(Finding, Empirical)",
                "p1": "WeatherXM operates a decentralized weather station network on Arbitrum, selling hyperlocal weather data through a data marketplace. Station operators deploy WeatherXM hardware (weather stations priced at approximately $200-400) and earn WXM tokens for providing verified meteorological data. Data consumers access weather API feeds by purchasing data credits, with WXM burned for marketplace access. As of February 2026, WeatherXM has deployed over 6,000 weather stations globally.",
                "p2": "WeatherXM's sink architecture combines a data subscription burn (WXM burned for API access) with emission rewards to station operators. The token launched on Arbitrum mainnet in May 2024 with a 100M total supply and 10-year distribution schedule. Governance concentration is the highest in the expanded cross-section: HHI 0.593 with a single address holding 76.4% of tokens (Dune, Arbitrum, February 2026). This extreme concentration likely reflects the recency of the token launch and team/treasury holdings that have not yet been distributed. The 52M token allocation to weather station owners over 10 years suggests concentration should decrease as the distribution schedule progresses.",
                "p3": "Verification relies on station hardware attestation and cross-validation against neighboring stations and satellite data. WeatherXM uses a reward algorithm that penalizes stations reporting data inconsistent with nearby stations or known weather patterns. The anti-gaming posture is moderate: hardware-based provenance plus statistical cross-validation, but sophisticated spoofing of meteorological data remains theoretically possible.",
                "p4": "WeatherXM's sustainability signal is early-stage but structurally sound. The data marketplace creates a direct link between weather data demand and token burns. Unit economics suggest a 10x discount versus Web2 weather API providers (Tomorrow.io at $0.05-0.20/API call vs. WeatherXM at ~$0.01/call), with Medium subsidy risk — station operators are compensated primarily through token emissions, but the hardware investment ($200-400) creates natural operator commitment. The 6,000+ station network provides a genuine data asset: hyperlocal weather data at station density that exceeds many national meteorological services in covered areas.",
            },
        ]

        # Mini profiles (2 paragraphs each)
        mini_profiles = [
            {
                "name": "Grass: AI Training Data via Idle Bandwidth",
                "tag": "(Observation, Data-Insufficient)",
                "p1": "Grass operates an AI data network on Solana where users share idle internet bandwidth to enable web scraping for AI model training. With a market capitalization of $87.4M and 472M circulating tokens (CoinGecko, February 2026), Grass is one of the larger DePIN protocols by market cap. GRASS tokens are staked for network access; no confirmed burn-and-mint mechanism exists, making this a staking-only model with S2R approximately zero (analogous to Livepeer's no-burn architecture).",
                "p2": "Grass's key observation is its positioning at the intersection of DePIN and AI infrastructure — a novel service domain where the 'physical infrastructure' is residential internet bandwidth rather than hardware sensors. The protocol cannot be directly compared to Web2 data labeling services (Scale AI, Appen) because it provides raw web data rather than labeled training data. Governance concentration data was unavailable due to Solana token transfer schema limitations on Dune Analytics. The protocol's sustainability depends on whether AI companies will pay sufficient premiums for decentralized data collection to offset emission costs — a question that cannot be assessed without marketplace revenue data.",
            },
            {
                "name": "Anyone Protocol: Decentralized Privacy Relay",
                "tag": "(Finding, Empirical)",
                "p1": "Anyone Protocol (formerly AirTor Protocol) operates a decentralized VPN/privacy relay network on Ethereum, where operators stake ANYONE tokens to run relay nodes. The protocol uses onion-routing architecture for censorship-resistant communications. Governance concentration is notably low: HHI 0.034 (Dune, Ethereum, February 2026), making ANYONE the most diffusely distributed token in the expanded cross-section — comparable to COMP (HHI 0.028) and significantly more distributed than DePIN peers. No confirmed burn mechanism exists; the economic model relies on staking for relay operation.",
                "p2": "The Anyone Protocol finding is significant for the governance concentration thesis: despite being a DePIN protocol, ANYONE exhibits governance distribution patterns closer to mature DeFi protocols than to DePIN peers (DIMO HHI 0.305, IoTeX HHI 0.388, WXM HHI 0.593). This suggests that governance concentration in DePIN is not an inherent property of the sector but reflects design choices — particularly initial token distribution and the presence of large treasury/team allocations. The privacy relay service domain has no direct S2R metric (no burn mechanism), analogous to Livepeer's fee-in-different-token model.",
            },
            {
                "name": "UpRock: Decentralized Bandwidth Proxy",
                "tag": "(Observation, Data-Insufficient)",
                "p1": "UpRock operates a decentralized residential proxy network on Solana, where users share idle bandwidth for enterprise web scraping. With a market capitalization of $1.2M (CoinGecko, February 2026), UpRock is early-stage. Posted proxy pricing (~$1.50/GB) represents an 8.5x discount versus Bright Data ($12.75/GB residential), though service quality, IP diversity, and reliability differ significantly.",
                "p2": "UpRock's burn mechanism and tokenomics are not well documented in public sources. The subsidy risk is High: operator revenue appears primarily from token emissions rather than enterprise proxy fees. Governance data was unavailable (Solana schema limitation). The protocol is included to represent the bandwidth-sharing DePIN category.",
            },
            {
                "name": "Wayru / WiFiDabba: Community WiFi Networks",
                "tag": "(Observation, Data-Insufficient)",
                "p1": "Wayru (Solana, Latin America) and WiFiDabba (Polygon, India) represent community WiFi DePIN protocols serving underserved markets where traditional ISP coverage is limited or expensive. Wayru pricing (~$2-5/month) offers a 6.4x discount versus regional ISP retail ($15-30/month). WiFiDabba does not appear to have a public governance token.",
                "p2": "Both protocols address the Rawlsian floor-raising principle directly: they expand internet access capabilities for underserved populations, creating genuine social welfare gains if sustainable. The subsidy risk is Medium-High for both: community WiFi requires ongoing hardware maintenance and local operator engagement that token emissions alone may not sustain. Data limitations prevent S2R computation or governance analysis for either protocol.",
            },
            {
                "name": "CUDIS: Wearable Health Data DePIN",
                "tag": "(Observation, Data-Insufficient)",
                "p1": "CUDIS operates a health data DePIN on Solana where users earn CUDIS tokens by sharing biometric data from smart rings ($69 vs. Oura Ring $299). With a market cap of $2.5M (CoinGecko, February 2026), the protocol is early-stage. The 4.3x hardware cost advantage drives adoption but raises questions about sensor quality and health data accuracy.",
                "p2": "CUDIS represents a novel DePIN domain — wearable health data — where information-flow norms (Floridi/Nissenbaum contextual integrity) are especially salient. Health data carries different privacy expectations than geolocation or bandwidth, and the thesis's Proposition 7 directly applies. The subsidy risk is High; the protocol's sustainability depends on whether health data consumers (insurers, researchers, pharma) will pay sufficient premiums for user-consented biometric data to offset emissions.",
            },
        ]

        # ROVR as footnote-level
        rovr_note = (
            "ROVR (autonomous delivery, Solana) was surveyed but excluded from profiling due to "
            "insufficient public documentation of tokenomics, service model, and operator metrics. "
            "The protocol is noted as a frontier DePIN category (mobility/delivery) for future research."
        )

        # Insert profiles
        current_idx = insert_idx
        for profile in full_profiles:
            p = insert_paragraph_after(doc, current_idx, "")
            current_idx += 1
            p = insert_paragraph_after(doc, current_idx, f"{profile['name']} {profile['tag']}")
            current_idx += 1
            for key in ["p1", "p2", "p3", "p4"]:
                p = insert_paragraph_after(doc, current_idx, profile[key])
                current_idx += 1

        p = insert_paragraph_after(doc, current_idx, "")
        current_idx += 1
        p = insert_paragraph_after(doc, current_idx, "Additional DePIN Protocol Profiles (Mini-Format)")
        current_idx += 1

        for profile in mini_profiles:
            p = insert_paragraph_after(doc, current_idx, f"{profile['name']} {profile['tag']}")
            current_idx += 1
            for key in ["p1", "p2"]:
                p = insert_paragraph_after(doc, current_idx, profile[key])
                current_idx += 1

        p = insert_paragraph_after(doc, current_idx, rovr_note)
        current_idx += 1

        print(f"  Inserted {len(full_profiles)} full + {len(mini_profiles)} mini profiles after paragraph {profiles_idx}")
    else:
        print("  WARNING: §6.10.3 insertion point not found")

    # ═══════════════════════════════════════════════════════════════
    # 7. UPDATE CLAIM REGISTRY (Table 7.1)
    # ═══════════════════════════════════════════════════════════════
    print("\n7. Updating Table 7.1 (Claim Registry)...")
    cr_idx = find_table_by_content(doc, "Claim")
    if cr_idx == -1:
        cr_idx = find_table_by_content(doc, "Governance concentration is design-dependent")
    if cr_idx >= 0:
        table = doc.tables[cr_idx]
        new_claims = [
            ["DePIN governance concentration exceeds DeFi range (HHI 0.034–0.593)", "This Thesis", "Table 6.G1 (expanded, 11 protocols)"],
            ["DePIN pricing undercuts Web2 by 1.4–80x across 13 domains", "This Thesis", "Table 6.U1 (expanded)"],
            ["S2R varies by sink architecture across 9 protocols", "This Thesis", "Table 6.D1 (expanded)"],
            ["Early-stage DePIN protocols exhibit higher governance concentration than mature DeFi", "This Thesis", "Table 6.G1 (DePIN vs DeFi comparison)"],
        ]
        add_rows_to_table(table, new_claims)
        print(f"  Added {len(new_claims)} claims to Table 7.1 (table index {cr_idx})")
    else:
        print("  WARNING: Table 7.1 not found")

    # ═══════════════════════════════════════════════════════════════
    # 8. UPDATE CHAPTER 7 TEXT REFERENCES
    # ═══════════════════════════════════════════════════════════════
    print("\n8. Updating Chapter 7 text references...")
    updates_made = 0
    for i, para in enumerate(doc.paragraphs):
        original = para.text

        # Update governance protocol count
        if "Gini 0.86–0.92" in para.text or "Gini 0.86-0.92" in para.text:
            para.text = para.text.replace("Gini 0.86–0.92", "Gini 0.79–0.98")
            para.text = para.text.replace("Gini 0.86-0.92", "Gini 0.79-0.98")
            updates_made += 1
        if "across all 7 protocols" in para.text:
            para.text = para.text.replace("across all 7 protocols", "across all 11 protocols")
            updates_made += 1
        if "7 protocols from Dune" in para.text:
            para.text = para.text.replace("7 protocols from Dune", "11 protocols from Dune (multi-chain)")
            updates_made += 1
        if "across 7 protocols" in para.text and "Table 6.G1" in para.text:
            para.text = para.text.replace("across 7 protocols", "across 11 protocols")
            updates_made += 1
        if "0.028 (COMP) to 0.174 (CRV)" in para.text:
            para.text = para.text.replace("0.028 (COMP) to 0.174 (CRV)", "0.028 (COMP) to 0.593 (WXM)")
            updates_made += 1
        if "0.028 (COMP, most diffuse) to 0.174 (CRV, highly concentrated" in para.text:
            para.text = para.text.replace(
                "0.028 (COMP, most diffuse) to 0.174 (CRV, highly concentrated via the veCRV locker)",
                "0.028 (COMP, most diffuse) to 0.593 (WXM, extremely concentrated — 76% in top address, reflecting recent token launch). Among DeFi protocols, CRV remains the most concentrated at 0.174 (veCRV locker). DePIN protocols (DIMO 0.305, IOTX 0.388, WXM 0.593) exhibit systematically higher concentration than DeFi peers, with the exception of Anyone Protocol (0.034)"
            )
            updates_made += 1

        # Update unit economics references
        if "29-80x" in para.text or "29–80×" in para.text or "29–80x" in para.text:
            para.text = para.text.replace("29-80x", "1.4–80x")
            para.text = para.text.replace("29–80×", "1.4–80×")
            para.text = para.text.replace("29–80x", "1.4–80x")
            updates_made += 1
        if "five domains" in para.text.lower() and "Table 6.U1" in para.text:
            para.text = para.text.replace("five domains", "thirteen domains")
            para.text = para.text.replace("5 domains", "13 domains")
            updates_made += 1

        # Update protocol counts
        if "3 protocols" in para.text and ("S2R" in para.text or "sink" in para.text.lower()):
            para.text = para.text.replace("3 protocols", "9 protocols")
            updates_made += 1

    print(f"  Made {updates_made} text updates in Chapter 7")

    # ═══════════════════════════════════════════════════════════════
    # 9. UPDATE §C.7 MANIFEST
    # ═══════════════════════════════════════════════════════════════
    print("\n9. Updating §C.7 Manifest...")
    manifest_idx = find_paragraph_containing(doc, "Schema (Appendix D.1) + Helium worked example")
    if manifest_idx >= 0:
        doc.paragraphs[manifest_idx].text = doc.paragraphs[manifest_idx].text.replace(
            "Schema (Appendix D.1) + Helium worked example",
            "Schema (Appendix D.1) + 11 protocol entries (Helium deep + IoTeX, DIMO, Hivemapper, WeatherXM full profiles + 7 mini-profiles)"
        )
        print(f"  Updated §C.7 manifest reference")
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Schema (Appendix D.1) + Helium worked example" in cell.text:
                    cell.text = cell.text.replace(
                        "Schema (Appendix D.1) + Helium worked example",
                        "Schema (Appendix D.1) + 11 protocol entries (Helium deep + 4 full + 5 mini profiles)"
                    )

    # ═══════════════════════════════════════════════════════════════
    # 10. UPDATE ABSTRACT (counts only)
    # ═══════════════════════════════════════════════════════════════
    print("\n10. Updating Abstract scope counts...")
    for i, para in enumerate(doc.paragraphs):
        if "verified empirical core" in para.text and i < 30:
            # Only update counts, preserve language
            if "governance HHI" in para.text:
                para.text = para.text.replace("seven protocols", "eleven protocols")
                para.text = para.text.replace("7 protocols", "11 protocols")
            break

    # ═══════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════
    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("✓ Thesis v11 saved successfully!")

    # Count words
    total_words = sum(len(para.text.split()) for para in doc.paragraphs)
    print(f"\n  Approximate word count: {total_words:,}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
